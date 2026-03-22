from __future__ import annotations

import re
from pathlib import Path


ARABIC_PATTERN = re.compile(r"[\u0600-\u06FF]")
MARKDOWN_BOLD_PATTERN = re.compile(r"(\*\*.*?\*\*)")
IMAGE_PATTERN = re.compile(r"!\[(.*?)\]\((.*?)\)")
PAGE_BREAK_PATTERN = re.compile(r"--- Page \d+ ---")
MONOSPACE_FONTS = {"consolas", "courier new", "menlo", "monaco", "source code pro"}


def _require_python_docx():
    try:
        from docx import Document
        from docx.document import Document as DocumentType
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.shared import Inches, Pt, RGBColor
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for document conversion. Install it from requirements.txt."
        ) from exc

    return {
        "Document": Document,
        "DocumentType": DocumentType,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "WD_TABLE_ALIGNMENT": WD_TABLE_ALIGNMENT,
        "OxmlElement": OxmlElement,
        "qn": qn,
        "Inches": Inches,
        "Pt": Pt,
        "RGBColor": RGBColor,
        "CT_P": CT_P,
        "CT_Tbl": CT_Tbl,
        "Paragraph": Paragraph,
        "Table": Table,
    }


def _normalize_text(value: str) -> str:
    return value.replace("\r\n", "\n").replace("\r", "\n")


def _detect_rtl(lines: list[str]) -> bool:
    return any(ARABIC_PATTERN.search(line) for line in lines[:20])


def _set_paragraph_style(paragraph, is_rtl: bool, qn, OxmlElement, alignment_enum) -> None:
    p_props = paragraph._p.get_or_add_pPr()

    for old_bidi in p_props.xpath("w:bidi"):
        p_props.remove(old_bidi)
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1" if is_rtl else "0")
    p_props.append(bidi)

    for old_jc in p_props.xpath("w:jc"):
        p_props.remove(old_jc)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right" if is_rtl else "left")
    p_props.append(jc)

    paragraph.alignment = alignment_enum.RIGHT if is_rtl else alignment_enum.LEFT


def _set_run_font(run, font_name: str, qn, OxmlElement) -> None:
    run.font.name = font_name
    r_props = run._r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:cs"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_props.append(r_fonts)


def _set_run_direction(run, is_rtl: bool, qn, OxmlElement) -> None:
    r_props = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1" if is_rtl else "0")
    r_props.append(rtl)


def _set_table_rtl(table, qn, OxmlElement) -> None:
    tbl_props = table._tbl.tblPr
    bidi_visual = OxmlElement("w:bidiVisual")
    bidi_visual.set(qn("w:val"), "1")
    tbl_props.append(bidi_visual)


def _set_cell_background(cell, color_hex: str, qn, OxmlElement) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_vertical_alignment(cell, align: str, qn, OxmlElement) -> None:
    tc_props = cell._tc.get_or_add_tcPr()
    vertical_align = OxmlElement("w:vAlign")
    vertical_align.set(qn("w:val"), align)
    tc_props.append(vertical_align)


def _set_cell_borders(cell, qn, OxmlElement, **kwargs) -> None:
    tc_props = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = OxmlElement(f"w:{edge}")
        for key, value in kwargs[edge].items():
            tag.set(qn(f"w:{key}"), str(value))
        tc_borders.append(tag)
    tc_props.append(tc_borders)


def _set_keep_with_next(paragraph, qn, OxmlElement) -> None:
    p_props = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    keep_next.set(qn("w:val"), "1")
    p_props.append(keep_next)


def _set_repeat_header(row, qn, OxmlElement) -> None:
    tr_props = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "1")
    tr_props.append(header)


def _set_code_block_style(paragraph, qn, OxmlElement, alignment_enum) -> None:
    _set_paragraph_style(paragraph, False, qn, OxmlElement, alignment_enum)
    p_props = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for border_name in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{border_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "4")
        edge.set(qn("w:space"), "4")
        edge.set(qn("w:color"), "D0D0D0")
        borders.append(edge)
    p_props.append(borders)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F4F4F4")
    p_props.append(shading)


def _add_code_run(paragraph, text: str, qn, OxmlElement, Pt, RGBColor, is_command: bool = False) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    if is_command:
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)
    r_props = run._r.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Consolas")
    r_fonts.set(qn("w:hAnsi"), "Consolas")
    r_props.append(r_fonts)
    _set_run_direction(run, False, qn, OxmlElement)


def _append_inline_markdown(paragraph, text: str, *, is_rtl: bool, font_name: str, qn, OxmlElement, alignment_enum) -> None:
    _set_paragraph_style(paragraph, is_rtl, qn, OxmlElement, alignment_enum)
    for part in MARKDOWN_BOLD_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) >= 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(part)
        _set_run_font(run, font_name, qn, OxmlElement)
        _set_run_direction(run, is_rtl, qn, OxmlElement)


def convert_markdown_to_docx(
    markdown_path: str | Path,
    output_path: str | Path,
    *,
    layout_mode: str = "auto",
    font_name: str = "Dubai",
    image_width_inches: float = 5.4,
    log_cb=None,
    progress_cb=None,
) -> dict[str, object]:
    docx_api = _require_python_docx()
    Document = docx_api["Document"]
    WD_ALIGN_PARAGRAPH = docx_api["WD_ALIGN_PARAGRAPH"]
    WD_TABLE_ALIGNMENT = docx_api["WD_TABLE_ALIGNMENT"]
    OxmlElement = docx_api["OxmlElement"]
    qn = docx_api["qn"]
    Inches = docx_api["Inches"]
    Pt = docx_api["Pt"]
    RGBColor = docx_api["RGBColor"]

    markdown_path = Path(markdown_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not markdown_path.exists():
        raise FileNotFoundError(f"Markdown file not found: {markdown_path}")

    raw_text = markdown_path.read_text(encoding="utf-8")
    lines = _normalize_text(raw_text).splitlines()
    is_rtl = True if layout_mode == "rtl" else False if layout_mode == "ltr" else _detect_rtl(lines)

    if progress_cb is not None:
        progress_cb(0.05)
    if log_cb is not None:
        log_cb(f"Loaded markdown file: {markdown_path}")

    document = Document()
    table_data: list[list[str]] = []
    in_table = False
    in_code_block = False
    stats = {
        "paragraphs": 0,
        "headings": 0,
        "tables": 0,
        "images": 0,
        "code_blocks": 0,
    }

    def flush_table() -> None:
        nonlocal in_table, table_data
        if not in_table or not table_data:
            in_table = False
            table_data = []
            return
        rows_count = len(table_data)
        cols_count = max((len(row) for row in table_data), default=0)
        if rows_count == 0 or cols_count == 0:
            in_table = False
            table_data = []
            return

        table = document.add_table(rows=rows_count, cols=cols_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT if is_rtl else WD_TABLE_ALIGNMENT.LEFT
        if is_rtl:
            _set_table_rtl(table, qn, OxmlElement)
        _set_repeat_header(table.rows[0], qn, OxmlElement)

        for row_index, row_data in enumerate(table_data):
            cells = table.rows[row_index].cells
            for col_index, cell_text in enumerate(row_data):
                if col_index >= len(cells):
                    continue
                cell = cells[col_index]
                paragraph = cell.paragraphs[0]
                _set_paragraph_style(paragraph, is_rtl, qn, OxmlElement, WD_ALIGN_PARAGRAPH)
                _set_cell_vertical_alignment(cell, "center", qn, OxmlElement)
                if row_index <= 1:
                    _set_keep_with_next(paragraph, qn, OxmlElement)

                border_style = {"sz": 4, "val": "single", "color": "DDDDDD"}
                cell_borders = {
                    "top": border_style,
                    "left": border_style,
                    "bottom": border_style,
                    "right": border_style,
                }
                if row_index == 0:
                    _set_cell_background(cell, "E1F5FE", qn, OxmlElement)
                    cell_borders["bottom"] = {"sz": 12, "val": "single", "color": "B0BEC5"}
                elif row_index % 2 == 0:
                    _set_cell_background(cell, "FAFAFA", qn, OxmlElement)
                _set_cell_borders(cell, qn, OxmlElement, **cell_borders)

                cleaned_text = cell_text.strip()
                if cleaned_text.startswith("**") and cleaned_text.endswith("**") and len(cleaned_text) >= 4:
                    run = paragraph.add_run(cleaned_text[2:-2])
                    run.bold = True
                else:
                    run = paragraph.add_run(cleaned_text)
                    if row_index == 0:
                        run.bold = True
                _set_run_font(run, font_name, qn, OxmlElement)
                _set_run_direction(run, is_rtl, qn, OxmlElement)

        stats["tables"] += 1
        table_data = []
        in_table = False

    total_lines = max(len(lines), 1)
    for index, original_line in enumerate(lines, start=1):
        raw_line = original_line.rstrip("\n")
        stripped = raw_line.strip()

        if progress_cb is not None:
            progress_cb(0.05 + (index / total_lines) * 0.9)

        image_match = IMAGE_PATTERN.search(stripped)
        if image_match:
            candidate = Path(image_match.group(2))
            resolved_image = candidate if candidate.is_absolute() else (markdown_path.parent / candidate).resolve()
            if resolved_image.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(resolved_image), width=Inches(image_width_inches))
                stats["images"] += 1
                continue
            if log_cb is not None:
                log_cb(f"Skipped missing image: {candidate}")

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                stats["code_blocks"] += 1
            continue

        if in_code_block:
            paragraph = document.add_paragraph()
            _set_code_block_style(paragraph, qn, OxmlElement, WD_ALIGN_PARAGRAPH)
            is_command = stripped.startswith("$ ") or stripped.startswith("# ")
            _add_code_run(paragraph, raw_line, qn, OxmlElement, Pt, RGBColor, is_command=is_command)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            continue

        cleaned = PAGE_BREAK_PATTERN.sub("", stripped)
        cleaned = re.sub(r"<\|ref\|>.*?<\|/ref\|>", "", cleaned)
        cleaned = re.sub(r"<\|det\|>.*?<\|/det\|>", "", cleaned)

        if cleaned.startswith("|") and cleaned.endswith("|"):
            in_table = True
            if "---" not in cleaned:
                table_data.append([cell.strip() for cell in cleaned.strip("|").split("|")])
            continue
        if in_table:
            flush_table()

        if not cleaned:
            continue
        if cleaned.startswith("<div") or cleaned.startswith("</div") or cleaned.startswith("<br"):
            continue
        if cleaned == "<!-- pagebreak -->":
            document.add_page_break()
            continue
        if cleaned == "---":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_style(paragraph, is_rtl, qn, OxmlElement, WD_ALIGN_PARAGRAPH)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("________________________________________")
            run.bold = True
            _set_run_font(run, font_name, qn, OxmlElement)
            _set_run_direction(run, is_rtl, qn, OxmlElement)
            continue

        if cleaned.startswith("#"):
            level = len(cleaned) - len(cleaned.lstrip("#"))
            title_text = cleaned.lstrip("#").strip().replace("**", "")
            paragraph = document.add_heading("", level=min(level, 9))
            _set_paragraph_style(paragraph, is_rtl, qn, OxmlElement, WD_ALIGN_PARAGRAPH)
            run = paragraph.add_run(title_text)
            _set_run_font(run, font_name, qn, OxmlElement)
            _set_run_direction(run, is_rtl, qn, OxmlElement)
            stats["headings"] += 1
            continue

        if cleaned.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _append_inline_markdown(
                paragraph,
                cleaned[2:].strip(),
                is_rtl=is_rtl,
                font_name=font_name,
                qn=qn,
                OxmlElement=OxmlElement,
                alignment_enum=WD_ALIGN_PARAGRAPH,
            )
            stats["paragraphs"] += 1
            continue

        paragraph = document.add_paragraph()
        _append_inline_markdown(
            paragraph,
            cleaned,
            is_rtl=is_rtl,
            font_name=font_name,
            qn=qn,
            OxmlElement=OxmlElement,
            alignment_enum=WD_ALIGN_PARAGRAPH,
        )
        stats["paragraphs"] += 1

    flush_table()
    document.save(output_path)
    if progress_cb is not None:
        progress_cb(1.0)
    if log_cb is not None:
        log_cb(f"Saved DOCX file: {output_path}")

    return {
        "input_path": str(markdown_path),
        "output_path": str(output_path),
        "layout_mode": layout_mode,
        "rtl": is_rtl,
        **stats,
    }


def _escape_markdown_cell(value: str) -> str:
    return value.replace("\\", "\\\\").replace("|", "\\|").replace("\n", "<br>")


def _is_page_break_paragraph(paragraph) -> bool:
    xml = paragraph._element.xml
    return "<w:br" in xml and 'w:type="page"' in xml


def _is_code_block_paragraph(paragraph) -> bool:
    if "<w:pBdr>" in paragraph._element.xml:
        return True
    font_names = {
        (run.font.name or "").strip().lower()
        for run in paragraph.runs
        if run.text and run.text.strip()
    }
    return bool(font_names) and font_names.issubset(MONOSPACE_FONTS)


def _format_markdown_run(text: str, *, bold: bool, italic: bool) -> str:
    if not text:
        return ""
    safe = text.replace("\r\n", "\n").replace("\r", "\n")
    if bold and italic:
        return f"***{safe}***"
    if bold:
        return f"**{safe}**"
    if italic:
        return f"*{safe}*"
    return safe


def _paragraph_text_to_markdown(paragraph) -> str:
    if _is_code_block_paragraph(paragraph):
        code_text = paragraph.text.rstrip()
        return f"```\n{code_text}\n```" if code_text else ""

    parts: list[str] = []
    for run in paragraph.runs:
        if not run.text:
            continue
        parts.append(_format_markdown_run(run.text, bold=bool(run.bold), italic=bool(run.italic)))
    return "".join(parts).strip()


def _extract_paragraph_images(
    paragraph,
    *,
    document,
    image_dir: Path | None,
    saved_images: dict[str, str],
) -> list[str]:
    if image_dir is None:
        return []

    refs: list[str] = []
    for node in paragraph._element.iter():
        if not str(node.tag).endswith("blip"):
            continue
        rel_id = node.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not rel_id:
            continue
        relative_ref = saved_images.get(rel_id)
        if relative_ref is None:
            image_part = document.part.related_parts.get(rel_id)
            if image_part is None:
                continue
            extension = image_part.content_type.split("/")[-1].replace("jpeg", "jpg")
            filename = f"image_{len(saved_images) + 1}.{extension}"
            target = image_dir / filename
            target.write_bytes(image_part.blob)
            relative_ref = image_dir.name + "/" + filename
            saved_images[rel_id] = relative_ref
        refs.append(f"![]({relative_ref})")
    return refs


def _paragraph_to_markdown_block(
    paragraph,
    *,
    document,
    image_dir: Path | None,
    saved_images: dict[str, str],
) -> str:
    blocks: list[str] = []
    if _is_page_break_paragraph(paragraph):
        blocks.append("<!-- pagebreak -->")

    text = _paragraph_text_to_markdown(paragraph)
    style_name = ""
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except Exception:
        style_name = ""

    if text:
        if style_name.startswith("heading "):
            try:
                level = max(1, min(6, int(style_name.split()[-1])))
            except Exception:
                level = 1
            blocks.append(f"{'#' * level} {text}")
        elif "list bullet" in style_name:
            blocks.append(f"- {text}")
        elif "list number" in style_name:
            blocks.append(f"1. {text}")
        else:
            blocks.append(text)

    blocks.extend(
        _extract_paragraph_images(
            paragraph,
            document=document,
            image_dir=image_dir,
            saved_images=saved_images,
        )
    )
    return "\n".join(blocks).strip()


def _table_to_markdown(table) -> str:
    rows = []
    for row in table.rows:
        rows.append([_escape_markdown_cell(cell.text.strip()) for cell in row.cells])
    if not rows:
        return ""

    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    header = normalized[0]
    separator = ["---"] * column_count
    body_rows = normalized[1:]

    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
    ]
    for row in body_rows:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _iter_document_blocks(document, Paragraph, Table, CT_P, CT_Tbl):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def convert_docx_to_markdown(
    docx_path: str | Path,
    output_path: str | Path,
    *,
    extract_images: bool = True,
    log_cb=None,
    progress_cb=None,
) -> dict[str, object]:
    docx_api = _require_python_docx()
    Document = docx_api["Document"]
    DocumentType = docx_api["DocumentType"]
    Paragraph = docx_api["Paragraph"]
    Table = docx_api["Table"]
    CT_P = docx_api["CT_P"]
    CT_Tbl = docx_api["CT_Tbl"]

    docx_path = Path(docx_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")

    document = Document(str(docx_path))
    if not isinstance(document, DocumentType):
        raise RuntimeError(f"Could not open DOCX document: {docx_path}")

    image_dir = output_path.parent / f"{output_path.stem}_media" if extract_images else None
    if image_dir is not None:
        image_dir.mkdir(parents=True, exist_ok=True)

    blocks = list(_iter_document_blocks(document, Paragraph, Table, CT_P, CT_Tbl))
    total_blocks = max(len(blocks), 1)
    saved_images: dict[str, str] = {}
    markdown_blocks: list[str] = []
    stats = {
        "paragraphs": 0,
        "headings": 0,
        "tables": 0,
        "images": 0,
        "code_blocks": 0,
    }

    if progress_cb is not None:
        progress_cb(0.05)
    if log_cb is not None:
        log_cb(f"Loaded DOCX file: {docx_path}")

    for index, block in enumerate(blocks, start=1):
        if progress_cb is not None:
            progress_cb(0.05 + (index / total_blocks) * 0.9)

        if isinstance(block, Paragraph):
            chunk = _paragraph_to_markdown_block(
                block,
                document=document,
                image_dir=image_dir,
                saved_images=saved_images,
            )
            if not chunk:
                continue
            markdown_blocks.append(chunk)
            stats["images"] = len(saved_images)
            if chunk.startswith("```"):
                stats["code_blocks"] += 1
            elif chunk.startswith("#"):
                stats["headings"] += 1
            else:
                stats["paragraphs"] += 1
        elif isinstance(block, Table):
            chunk = _table_to_markdown(block)
            if not chunk:
                continue
            markdown_blocks.append(chunk)
            stats["tables"] += 1

    markdown_text = "\n\n".join(block for block in markdown_blocks if block).strip()
    output_path.write_text((markdown_text + "\n") if markdown_text else "", encoding="utf-8")

    if image_dir is not None and image_dir.exists() and not any(image_dir.iterdir()):
        try:
            image_dir.rmdir()
        except OSError:
            pass

    if progress_cb is not None:
        progress_cb(1.0)
    if log_cb is not None:
        log_cb(f"Saved markdown file: {output_path}")

    return {
        "input_path": str(docx_path),
        "output_path": str(output_path),
        "image_dir": str(image_dir) if image_dir is not None and image_dir.exists() and any(image_dir.iterdir()) else "",
        **stats,
    }
